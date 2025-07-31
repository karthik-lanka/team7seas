# import requests
# import fitz  # PyMuPDF
# from docx import Document
# import tempfile
# import os
# from urllib.parse import urlparse
# import logging

# logger = logging.getLogger(__name__)

# class DocumentProcessor:
#     """Handles downloading and processing of PDF and DOCX documents"""
    
#     def __init__(self):
#         self.session = requests.Session()
#         self.session.headers.update({
#             'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
#         })
    
#     def process_document(self, document_url: str) -> str:
#         """
#         Download and extract text from a document URL
        
#         Args:
#             document_url: URL to the PDF or DOCX document
            
#         Returns:
#             Extracted text content
#         """
#         try:
#             # Download the document
#             logger.info(f"Downloading document from: {document_url}")
#             response = self.session.get(document_url, timeout=30)
#             response.raise_for_status()
            
#             # Determine file type from URL or content type
#             parsed_url = urlparse(document_url)
#             file_extension = os.path.splitext(parsed_url.path)[1].lower()
            
#             if not file_extension:
#                 content_type = response.headers.get('content-type', '').lower()
#                 if 'pdf' in content_type:
#                     file_extension = '.pdf'
#                 elif 'word' in content_type or 'officedocument' in content_type:
#                     file_extension = '.docx'
#                 else:
#                     # Try to detect from content
#                     content_start = response.content[:4]
#                     if content_start == b'%PDF':
#                         file_extension = '.pdf'
#                     elif content_start[:2] == b'PK':  # ZIP-based format like DOCX
#                         file_extension = '.docx'
#                     else:
#                         raise ValueError("Unable to determine document type")
            
#             # Process based on file type
#             if file_extension == '.pdf':
#                 return self._extract_pdf_text(response.content)
#             elif file_extension == '.docx':
#                 return self._extract_docx_text(response.content)
#             else:
#                 raise ValueError(f"Unsupported file type: {file_extension}")
                
#         except requests.RequestException as e:
#             logger.error(f"Error downloading document: {str(e)}")
#             raise Exception(f"Failed to download document: {str(e)}")
#         except Exception as e:
#             logger.error(f"Error processing document: {str(e)}")
#             raise Exception(f"Failed to process document: {str(e)}")
    
#     def _extract_pdf_text(self, pdf_content: bytes) -> str:
#         """Extract text from PDF content using PyMuPDF"""
#         try:
#             with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
#                 temp_file.write(pdf_content)
#                 temp_file.flush()
                
#                 # Open PDF with PyMuPDF
#                 doc = fitz.open(temp_file.name)
#                 text_content = []
                
#                 for page_num in range(doc.page_count):
#                     page = doc[page_num]
#                     text = page.get_text()
#                     if text.strip():
#                         text_content.append(text)
                
#                 doc.close()
#                 os.unlink(temp_file.name)
                
#                 full_text = '\n\n'.join(text_content)
#                 logger.info(f"Extracted {len(full_text)} characters from PDF")
#                 return full_text
                
#         except Exception as e:
#             logger.error(f"Error extracting PDF text: {str(e)}")
#             raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
#     def _extract_docx_text(self, docx_content: bytes) -> str:
#         """Extract text from DOCX content using python-docx"""
#         try:
#             with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
#                 temp_file.write(docx_content)
#                 temp_file.flush()
                
#                 # Open DOCX with python-docx
#                 doc = Document(temp_file.name)
#                 text_content = []
                
#                 for paragraph in doc.paragraphs:
#                     if paragraph.text.strip():
#                         text_content.append(paragraph.text)
                
#                 # Also extract text from tables
#                 for table in doc.tables:
#                     for row in table.rows:
#                         for cell in row.cells:
#                             if cell.text.strip():
#                                 text_content.append(cell.text)
                
#                 os.unlink(temp_file.name)
                
#                 full_text = '\n\n'.join(text_content)
#                 logger.info(f"Extracted {len(full_text)} characters from DOCX")
#                 return full_text
                
#         except Exception as e:
#             logger.error(f"Error extracting DOCX text: {str(e)}")
#             raise Exception(f"Failed to extract text from DOCX: {str(e)}")

import requests
import fitz  # PyMuPDF
from docx import Document
import tempfile
import os
from io import BytesIO
from urllib.parse import urlparse
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles downloading and processing of PDF and DOCX documents"""

    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'
        })

    def process_document(self, document_url: str) -> str:
        """
        Download and extract text from a document URL.
        """
        try:
            logger.info(f"Downloading document from: {document_url}")
            response = self.session.get(document_url, timeout=15)
            response.raise_for_status()

            file_ext = self._detect_file_type(document_url, response)
            content = response.content

            if file_ext == '.pdf':
                return self._extract_pdf_text(content)
            elif file_ext == '.docx':
                return self._extract_docx_text(content)
            else:
                raise ValueError(f"Unsupported document type: {file_ext}")

        except Exception as e:
            logger.error(f"Document processing error: {str(e)}")
            raise Exception(f"Failed to process document: {str(e)}")

    def _detect_file_type(self, url: str, response: requests.Response) -> str:
        """Try to determine the document file type by URL or response headers."""
        parsed = urlparse(url)
        ext = os.path.splitext(parsed.path)[1].lower()

        if ext in ['.pdf', '.docx']:
            return ext

        ctype = response.headers.get('content-type', '').lower()
        if 'pdf' in ctype:
            return '.pdf'
        elif 'word' in ctype or 'officedocument' in ctype:
            return '.docx'

        # Fallback to content sniffing
        start = response.content[:4]
        if start == b'%PDF':
            return '.pdf'
        elif start[:2] == b'PK':
            return '.docx'
        else:
            raise ValueError("Could not determine document type")

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF content using PyMuPDF (in-memory)."""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_chunks = [page.get_text() for page in doc if page.get_text().strip()]
            full_text = '\n\n'.join(text_chunks)
            logger.info(f"Extracted {len(full_text)} characters from PDF")
            return full_text
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise Exception("Failed to extract text from PDF")

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX content using python-docx (in-memory)."""
        try:
            buffer = BytesIO(content)
            doc = Document(buffer)
            text_content = []

            for p in doc.paragraphs:
                if p.text.strip():
                    text_content.append(p.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            full_text = '\n\n'.join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise Exception("Failed to extract text from DOCX")
